from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from io import BytesIO
import datetime

class DocumentGenerator:
    def __init__(self):
        """Initialize Word document generator"""
        pass
    
    def create_document(self, transcription, original_filename):
        """Create Word document with transcribed content"""
        try:
            # Create new document
            doc = Document()
            
            # Set default font for the document
            self._set_default_font(doc)
            
            # Add title
            title = doc.add_heading('手書きアンケート文字起こし結果', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._apply_font_formatting(title, 'MS Gothic', 14, True)
            
            # Add metadata
            doc.add_paragraph()
            meta_para = doc.add_paragraph()
            meta_run1 = meta_para.add_run('元ファイル名: ')
            meta_run1.bold = True
            meta_run2 = meta_para.add_run(original_filename)
            self._apply_font_formatting(meta_para, 'MS Gothic', 11)
            
            meta_para = doc.add_paragraph()
            meta_run1 = meta_para.add_run('処理日時: ')
            meta_run1.bold = True
            meta_run2 = meta_para.add_run(datetime.datetime.now().strftime('%Y年%m月%d日 %H:%M:%S'))
            self._apply_font_formatting(meta_para, 'MS Gothic', 11)
            
            meta_para = doc.add_paragraph()
            meta_run1 = meta_para.add_run('処理方法: ')
            meta_run1.bold = True
            meta_run2 = meta_para.add_run('Gemini AI による手書き文字認識')
            self._apply_font_formatting(meta_para, 'MS Gothic', 11)
            
            # Add separator
            sep_para = doc.add_paragraph('=' * 50)
            self._apply_font_formatting(sep_para, 'MS Gothic', 11)
            
            # Add content heading
            content_heading = doc.add_heading('文字起こし内容', level=1)
            self._apply_font_formatting(content_heading, 'MS Gothic', 12, True)
            
            # Process and add transcribed content
            self._add_transcription_content(doc, transcription)
            
            # Add footer note
            doc.add_paragraph()
            sep_para = doc.add_paragraph('=' * 50)
            self._apply_font_formatting(sep_para, 'MS Gothic', 11)
            
            footer = doc.add_paragraph()
            footer_run = footer.add_run('注意事項:')
            footer_run.bold = True
            self._apply_font_formatting(footer, 'MS Gothic', 11)
            
            footer = doc.add_paragraph(
                '• この文書は手書き文字の自動認識結果です。\n'
                '• 認識精度には限界があり、誤認識の可能性があります。\n'
                '• 重要な内容については原本をご確認ください。'
            )
            self._apply_font_formatting(footer, 'MS Gothic', 11)
            
            # Save to BytesIO
            doc_buffer = BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            return doc_buffer.getvalue()
            
        except Exception as e:
            raise Exception(f"Word文書生成エラー: {str(e)}")
    
    def _add_transcription_content(self, doc, transcription):
        """Add transcribed content to document with proper formatting"""
        if not transcription or transcription.strip() == "":
            para = doc.add_paragraph("文字起こし結果がありません。")
            self._apply_font_formatting(para, 'MS Gothic', 11)
            return
        
        # Process content with question-based formatting
        self._add_formatted_text_with_questions(doc, transcription)
    
    def _add_formatted_text(self, doc, text):
        """Add text with proper formatting, preserving structure"""
        paragraphs = text.split('\n\n')
        
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue
            
            # Check if it's a section header (contains specific patterns)
            if self._is_section_header(para_text):
                heading = doc.add_heading(para_text, level=3)
                self._apply_font_formatting(heading, 'ＭＳ ゴシック', 11, True)
            else:
                # Regular paragraph
                lines = para_text.split('\n')
                if len(lines) == 1:
                    # Single line paragraph
                    para = doc.add_paragraph(para_text)
                    self._apply_font_formatting(para, 'MS Gothic', 11)
                else:
                    # Multi-line paragraph - preserve line breaks
                    para = doc.add_paragraph()
                    for i, line in enumerate(lines):
                        line = line.strip()
                        if line:
                            if i > 0:
                                para.add_run('\n')
                            para.add_run(line)
                    self._apply_font_formatting(para, 'MS Gothic', 11)
    
    def _is_section_header(self, text):
        """Determine if text should be formatted as a section header"""
        # Simple heuristics for Japanese section headers
        header_indicators = [
            '質問', '問', 'Q', 'A', '回答', '設問', '項目', 
            'アンケート', '調査', '意見', 'コメント'
        ]
        
        # Check if text is short and contains header indicators
        if len(text) < 50:
            for indicator in header_indicators:
                if indicator in text:
                    return True
        
        # Check if text ends with common header endings
        header_endings = ['：', ':', '）', ')', '？', '?']
        if any(text.endswith(ending) for ending in header_endings):
            return True
        
        return False
    
    def _add_formatted_text_with_questions(self, doc, text):
        """Add text with question-based formatting and proper spacing"""
        # Split by lines to process question-answer pairs
        lines = text.split('\n')
        current_question = None
        current_answers = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Check if line is a question (starts with Q, 問, 質問, etc.)
            if self._is_question_line(line):
                # Add previous question-answer pair if exists
                if current_question:
                    self._add_question_answer_pair(doc, current_question, current_answers)
                
                current_question = line
                current_answers = []
            elif line.startswith('・'):
                # This is a bullet point answer
                if current_question:
                    answer_text = line[1:].strip()  # Remove bullet point
                    current_answers.append(answer_text)
                else:
                    # No question context, add as regular paragraph
                    para = doc.add_paragraph(line)
                    self._apply_font_formatting(para, 'MS Gothic', 11)
            else:
                # This is likely an answer or continuation
                if current_question:
                    current_answers.append(line)
                else:
                    # No question context, add as regular paragraph
                    para = doc.add_paragraph(line)
                    self._apply_font_formatting(para, 'MS Gothic', 11)
        
        # Add the last question-answer pair
        if current_question:
            self._add_question_answer_pair(doc, current_question, current_answers)
    
    def _is_question_line(self, line):
        """Check if a line is a question"""
        question_indicators = ['Q', 'Ｑ', '問', '質問', 'アンケート']
        question_endings = ['？', '?', '：', ':']
        
        # Check for question indicators at the start
        for indicator in question_indicators:
            if line.startswith(indicator):
                return True
        
        # Check for question endings
        for ending in question_endings:
            if line.endswith(ending):
                return True
                
        return False
    
    def _add_question_answer_pair(self, doc, question, answers):
        """Add a question and its answers with proper formatting"""
        # Add question with proper spacing and bold formatting
        q_para = doc.add_paragraph()
        q_run = q_para.add_run(question)
        q_run.bold = True
        self._apply_font_formatting(q_para, 'MS Gothic', 11, bold=True)
        
        # Add answers as bullet points (with line break after question)
        if answers:
            for answer in answers:
                if answer.strip():  # Only add non-empty answers
                    a_para = doc.add_paragraph(f"・{answer.strip()}")
                    self._apply_font_formatting(a_para, 'MS Gothic', 11)
        
        # Add spacing after each question-answer pair
        doc.add_paragraph()
    
    def _set_default_font(self, doc):
        """Set default font for the document using multiple methods"""
        try:
            # Method 1: Set Normal style font
            style = doc.styles['Normal']
            font = style.font
            font.name = 'MS Gothic'
            font.size = Pt(11)
            
            # Method 2: Set document defaults via XML
            from docx.oxml import OxmlElement
            
            # Get document defaults
            doc_defaults = doc.element.xpath('//w:docDefaults')[0]
            char_props = doc_defaults.xpath('.//w:rPrDefault/w:rPr')[0]
            
            # Remove existing font elements
            for fonts_elem in char_props.xpath('.//w:rFonts'):
                char_props.remove(fonts_elem)
            
            # Create new font element with MS Gothic
            fonts = OxmlElement('w:rFonts')
            fonts.set(qn('w:ascii'), 'MS Gothic')
            fonts.set(qn('w:eastAsia'), 'MS Gothic') 
            fonts.set(qn('w:hAnsi'), 'MS Gothic')
            fonts.set(qn('w:cs'), 'MS Gothic')
            char_props.insert(0, fonts)
            
            # Method 3: Set theme fonts
            theme_part = doc.part.package.part_related_by(
                'http://schemas.openxmlformats.org/officeDocument/2006/relationships/theme'
            )
            if theme_part:
                theme_element = theme_part.element
                font_scheme = theme_element.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}fontScheme')
                if font_scheme is not None:
                    # Update major font (headings)
                    major_font = font_scheme.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}majorFont')
                    if major_font is not None:
                        latin = major_font.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}latin')
                        if latin is not None:
                            latin.set('typeface', 'MS Gothic')
                    
                    # Update minor font (body text)
                    minor_font = font_scheme.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}minorFont')
                    if minor_font is not None:
                        latin = minor_font.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}latin')
                        if latin is not None:
                            latin.set('typeface', 'MS Gothic')
                            
        except Exception as e:
            # Fallback method
            try:
                style = doc.styles['Normal']
                style.font.name = 'MS Gothic'
                style.font.size = Pt(11)
            except:
                pass
    
    def _apply_font_formatting(self, paragraph, font_name='MS Gothic', font_size=11, bold=False):
        """Apply font formatting to paragraph with strong enforcement"""
        try:
            # If paragraph has no text, add empty run first
            if not paragraph.text.strip() and not paragraph.runs:
                paragraph.add_run("")
            
            # Apply font to all existing runs
            for run in paragraph.runs:
                run.font.name = font_name
                run.font.size = Pt(font_size)
                if bold:
                    run.bold = True
                    
                # Force font family at XML level for each run
                from docx.oxml import OxmlElement
                rpr = run._element.get_or_add_rPr()
                
                # Remove existing font elements
                for fonts_elem in rpr.xpath('.//w:rFonts'):
                    rpr.remove(fonts_elem)
                
                # Add new font element
                fonts = OxmlElement('w:rFonts')
                fonts.set(qn('w:ascii'), font_name)
                fonts.set(qn('w:eastAsia'), font_name)
                fonts.set(qn('w:hAnsi'), font_name)
                fonts.set(qn('w:cs'), font_name)
                rpr.insert(0, fonts)
            
            # Set paragraph style font
            try:
                paragraph.style.font.name = font_name
                paragraph.style.font.size = Pt(font_size)
            except:
                pass
                
        except Exception:
            # Final fallback
            try:
                for run in paragraph.runs:
                    run.font.name = font_name
                    run.font.size = Pt(font_size)
                    if bold:
                        run.bold = True
            except:
                pass
